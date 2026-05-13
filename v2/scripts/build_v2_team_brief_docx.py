"""Build the v2 end-to-end team brief DOCX.

This script intentionally keeps the document generation deterministic. The
brief is not a final research report; it is the handoff document a team lead
can send to five teammates before using the limited server allocation.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "v2_end_to_end_team_brief.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RISK = "9B1C1C"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Pt(widths[index] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 120, line: int = 264) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def _add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("HateSpeachStudy v2 team brief")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def _add_title(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    title.add_run("v2 End-to-End 모델 및 업무하달 브리프")
    subtitle = document.add_paragraph()
    subtitle.add_run("목적: 팀원 5명이 v2만 보고 학습, 통계, XAI, report/dashboard 검증을 나눠 수행하도록 기준을 고정한다.").bold = True
    meta = document.add_paragraph()
    meta.add_run("Run ID: ").bold = True
    meta.add_run("v2_15seed   ")
    meta.add_run("Canonical output: ").bold = True
    meta.add_run("v2/outputs/experiments/v2_15seed/   ")
    meta.add_run("작성 기준: ").bold = True
    meta.add_run("2026-05-13")


def _add_callout(document: Document, title: str, body: str, color: str = CALLOUT) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_width(table, [9360])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, color)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE if color != RISK else "FFFFFF")
    paragraph.add_run(f"\n{body}")


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        _set_paragraph_spacing(paragraph, after=80, line=280)


def _add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)
        _set_paragraph_spacing(paragraph, after=80, line=280)


def _add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        _set_cell_shading(header_cells[index], LIGHT_GRAY)
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    _set_table_width(table, widths)
    document.add_paragraph()


def build() -> Path:
    document = Document()
    _style_document(document)
    _add_footer(document)
    _add_title(document)

    _add_callout(
        document,
        "결론",
        "v2는 독립 실행 기준으로 정리됐다. 팀원은 v1을 열지 않고 v2/runtime, v2/pipeline, v2/docs만 기준으로 학습 smoke, 통계 집계, XAI evidence bundle, report/dashboard 검증을 진행한다.",
    )

    document.add_heading("1. v2의 기준 원칙", level=1)
    _add_bullets(
        document,
        [
            "v1은 archive/reference이며 새 실행의 기준 코드가 아니다.",
            "학습, 평가, 추론, XAI runtime은 v2/runtime 안에 둔다.",
            "end-to-end orchestration은 v2/pipeline 안에서 관리한다.",
            "모든 새 산출물은 v2/outputs/experiments/v2_15seed/ 아래에 저장한다.",
            "full benchmark는 smoke gate 통과 전에는 시작하지 않는다.",
        ],
    )

    document.add_heading("2. End-to-End 파이프라인", level=1)
    _add_numbered(
        document,
        [
            "benchmark: 8조건 x 15 seed 학습 실행과 metrics/history/predictions/checkpoint 저장",
            "aggregate: 완료 run의 metrics를 모아 summary, paired tests, Holm 보정 결과 생성",
            "xai-primary: A_B와 D_B를 모든 seed에서 비교해 explanation stability 확인",
            "xai-deep: median-performing checkpoint로 상세 case 분석",
            "xai-ablation: 8조건 전체의 경량 XAI 비교",
            "xai-bundle: raw XAI artifact를 report/dashboard용 evidence bundle로 통합",
            "report: final_report.md와 final_report.docx 생성",
            "dashboard: dashboard/index.html 생성",
        ],
    )

    document.add_heading("3. 현재 코드 상태", level=1)
    _add_table(
        document,
        ["영역", "상태", "다음 확인"],
        [
            ["v2 runtime", "v2/runtime에 학습/XAI/대시보드 코드 배치", "A_B seed 42 실제 학습 smoke"],
            ["benchmark adapter", "training_adapter.py가 v2/runtime 학습 호출", "metrics, predictions, checkpoint 생성 확인"],
            ["statistics", "paired test, Holm, CI, Cohen dz 계산 가능", "smoke 결과로 paired row 생성 확인"],
            ["XAI bundle", "evidence bundle 파일 계약 생성 가능", "실제 XAI artifact를 읽어 claim 채우기"],
            ["report/dashboard", "Markdown, Word, HTML scaffold 생성", "실제 결과 기반 표/문장 렌더링"],
        ],
        [1700, 3560, 4100],
    )

    document.add_heading("4. 팀원 5명 역할 배분", level=1)
    _add_table(
        document,
        ["사람", "역할", "기간", "책임 코드"],
        [
            ["1번", "Runtime Training 검증", "D0-D2", "v2/runtime/experiment_core.py, utils.py"],
            ["2번", "Adapter/CLI 검증", "D0-D3", "v2/pipeline/training_adapter.py, runner.py, artifacts.py"],
            ["3번", "Statistics / Inference Output", "D1-D6", "v2/pipeline/statistics.py, schema.py"],
            ["4번", "XAI Runtime / Evidence Bundle", "D2-D8", "v2/runtime/experiment_xai.py, xai.py, xai_bundle.py"],
            ["5번", "Integration / Report / Server", "D0-D10", "v2/run.sh, cli.py, reporting.py, dashboard code"],
        ],
        [900, 2100, 1100, 5260],
    )

    document.add_heading("5. 서버 실행 Gate", level=1)
    _add_callout(
        document,
        "Full run 금지 조건",
        "아래 smoke gate가 통과하기 전에는 8 conditions x 15 seeds = 120 benchmark를 시작하지 않는다.",
        color="FFF2CC",
    )
    _add_bullets(
        document,
        [
            "v2_runtime_import_smoke 통과",
            "A_B seed 42 단일 학습 성공",
            "A_B/D_B seed 42 paired smoke 성공",
            "metrics.json, history.csv, run_config.json, predictions.csv, checkpoint 생성",
            "aggregate가 smoke 결과를 읽고 paired_tests row 생성",
            "checkpoint_path와 predictions_path가 v2 output root 내부를 가리킴",
            "report/dashboard stage가 실패하지 않음",
        ],
    )

    document.add_heading("6. 실행 명령", level=1)
    _add_table(
        document,
        ["목적", "명령"],
        [
            ["상태 확인", "PYTHON_BIN=python3 ./v2/run.sh e2e status --run-id v2_15seed"],
            ["dry-run", "PYTHON_BIN=python3 ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B,D_B --seeds 42 --dry-run"],
            ["첫 smoke", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B --seeds 42 --execute --resume"],
            ["paired smoke", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B,D_B --seeds 42 --execute --resume"],
            ["집계", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e aggregate --run-id v2_15seed"],
            ["XAI bundle", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e xai-bundle --run-id v2_15seed"],
            ["보고서", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e report --run-id v2_15seed"],
        ],
        [1800, 7560],
    )

    document.add_heading("7. 통계와 XAI 해석 기준", level=1)
    _add_bullets(
        document,
        [
            "15 seed는 학습 stochasticity를 추정하기 위한 반복 단위다.",
            "조건 비교는 같은 seed끼리 묶는 paired design으로 계산한다.",
            "p-value만 보고하지 않고 mean difference, 95% CI, effect size를 함께 보고한다.",
            "여러 비교는 Holm-Bonferroni로 보정한다.",
            "XAI는 모델 설계의 인과 증명이 아니라 사후 검증 근거다.",
            "report/dashboard는 raw XAI case보다 xai_claims.json과 xai_dashboard_bundle.json을 우선 소비한다.",
        ],
    )

    document.add_heading("8. 팀원에게 보낼 기본 문장", level=1)
    _add_callout(
        document,
        "복사해서 전달",
        "우리는 HateSpeachStudy의 v2_15seed end-to-end 파이프라인을 구현 중입니다. 실행 기준 코드는 v2/runtime과 v2/pipeline입니다. 모든 새 산출물은 v2/outputs/experiments/v2_15seed/ 아래에 저장해야 합니다. 자기 담당 파일 밖을 수정해야 하면 이유와 변경 범위를 먼저 설명해주세요.",
    )

    document.add_heading("9. Claude/Gemini/Cursor용 Portable AI Skills", level=1)
    _add_bullets(
        document,
        [
            "Claude Code는 v2/CLAUDE.md를 먼저 읽고, 자기 역할에 맞는 v2/ai_skills/*/SKILL.md를 읽게 한다.",
            "Gemini CLI 또는 Gemini Code Assist는 v2/GEMINI.md와 v2/ai_skills/common_project_rules.md를 먼저 읽게 한다.",
            "Cursor, Antigravity, chat형 도구는 v2/ai_skills/README.md의 사용법과 역할별 SKILL.md를 그대로 붙여 넣어도 된다.",
            "Codex 계열 도구는 v2/ai_skills/hatespeech-v2-* 폴더를 local skills directory로 복사해도 된다.",
            "모든 도구는 같은 handoff 형식과 같은 smoke/full-run gate를 사용한다.",
        ],
    )

    _add_table(
        document,
        ["작업", "읽힐 Skill"],
        [
            ["전체 통합", "v2/ai_skills/hatespeech-v2-e2e/SKILL.md"],
            ["학습/benchmark", "v2/ai_skills/hatespeech-v2-benchmark/SKILL.md"],
            ["통계/aggregate", "v2/ai_skills/hatespeech-v2-statistics/SKILL.md"],
            ["XAI/evidence bundle", "v2/ai_skills/hatespeech-v2-xai/SKILL.md"],
            ["report/dashboard", "v2/ai_skills/hatespeech-v2-report-dashboard/SKILL.md"],
            ["리뷰/preflight", "v2/ai_skills/hatespeech-v2-review/SKILL.md"],
        ],
        [2600, 6760],
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
