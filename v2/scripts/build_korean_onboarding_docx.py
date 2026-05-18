"""Build a Korean onboarding DOCX for immediate v2 team work.

This document turns the reading index into an operational checklist: what each
teammate reads first, what they run next, and what they must report before the
limited server window is used.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "한글_필독문서_업무도입_가이드.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WARNING = "FFF2CC"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def _add_footer(document: Document) -> None:
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("HateSpeachStudy v2 Korean onboarding guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def _add_title(document: Document) -> None:
    title = document.add_paragraph(style="Title")
    title.add_run("한글 필독문서 및 업무도입 가이드")
    paragraph = document.add_paragraph()
    paragraph.add_run("대상: ").bold = True
    paragraph.add_run("v2_15seed 팀원 5명, AI 도구 사용자, 서버 실행 담당자")
    paragraph = document.add_paragraph()
    paragraph.add_run("목적: ").bold = True
    paragraph.add_run("문서 읽기에서 멈추지 않고, 역할별 첫 업무와 검증 명령까지 바로 이어지게 한다.")


def _add_callout(document: Document, title: str, body: str, fill: str = CALLOUT) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_width(table, [9360])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(f"\n{body}")


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def _add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def _add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        _set_cell_shading(cell, LIGHT_GRAY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    _set_table_width(table, widths)
    document.add_paragraph()


def build() -> Path:
    document = Document()
    _style_document(document)
    _add_footer(document)
    _add_title(document)

    _add_callout(
        document,
        "검사 결과",
        "17_korean_reading_file_index.md는 문서 안내용으로 충분하다. 다만 바로 업무에 투입하려면 역할별 첫 행동, 실행 금지 조건, 완료 보고 양식이 한 문서에 묶여 있어야 하므로 이 Word 가이드를 추가한다.",
    )

    document.add_heading("1. 업무 시작 전 20분 읽기", level=1)
    _add_numbered(
        document,
        [
            "v2/README.md에서 v2가 현재 기준 workspace임을 확인한다.",
            "v2/docs/17_korean_reading_file_index.md에서 자기 역할의 읽기 코스를 고른다.",
            "v2/docs/14_team_assignment_matrix.md에서 자기 담당 코드와 기간을 확인한다.",
            "v2/docs/15_runtime_code_validation_matrix.md에서 검증 질문과 명령을 확인한다.",
            "AI 도구를 쓴다면 v2/docs/16_portable_ai_agent_skills_guide.md와 역할별 SKILL.md를 읽힌다.",
        ],
    )

    document.add_heading("2. 역할별 바로 할 일", level=1)
    _add_table(
        document,
        ["역할", "먼저 읽을 문서", "첫 업무"],
        [
            ["1번 Runtime Training", "14, 15, experiment_core 관련 docs", "BASE_DIR/output path와 train_neural_model 산출물 확인"],
            ["2번 Adapter/CLI", "10, 15, benchmark agent", "A_B/D_B seed 42 dry-run과 v2_runtime_import_smoke 확인"],
            ["3번 Statistics", "03, 07, statistics agent", "aggregate가 빈/부분 결과에서도 CSV contract를 지키는지 확인"],
            ["4번 XAI", "04, 08, xai/evidence docs", "xai-bundle 산출물 계약과 claim source 연결 기준 확인"],
            ["5번 Integration", "06, 11, 15, review skill", "compile/config/dry-run/report/dashboard preflight 실행"],
        ],
        [1800, 3300, 4260],
    )

    document.add_heading("3. 서버 실행 전 금지선", level=1)
    _add_callout(
        document,
        "Full benchmark 금지",
        "A_B seed 42 단일 smoke와 A_B/D_B seed 42 paired smoke가 끝나기 전에는 full 120 benchmark를 시작하지 않는다.",
        fill=WARNING,
    )
    _add_bullets(
        document,
        [
            "v2_runtime_import_smoke가 통과해야 한다.",
            "metrics.json, history.csv, run_config.json, predictions.csv, checkpoint가 v2 output root 아래 생성되어야 한다.",
            "aggregate가 smoke 결과를 읽고 paired_tests row를 만들어야 한다.",
            "report/dashboard stage가 실패하지 않아야 한다.",
            "실패 run 재실행 계획 없이 full run을 다시 돌리지 않는다.",
        ],
    )

    document.add_heading("4. 팀 공통 검증 명령", level=1)
    _add_table(
        document,
        ["목적", "명령"],
        [
            ["문법 확인", "python3 -m compileall v2/runtime v2/pipeline v2/scripts"],
            ["config 확인", "python3 -m json.tool v2/configs/v2_15seed.json >/tmp/v2_config_check.json"],
            ["benchmark dry-run", "PYTHON_BIN=python3 ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B,D_B --seeds 42 --dry-run"],
            ["aggregate smoke", "PYTHON_BIN=python3 ./v2/run.sh e2e aggregate --run-id v2_15seed"],
            ["xai-bundle smoke", "PYTHON_BIN=python3 ./v2/run.sh e2e xai-bundle --run-id v2_15seed"],
            ["report/dashboard", "PYTHON_BIN=python3 ./v2/run.sh e2e report --run-id v2_15seed && PYTHON_BIN=python3 ./v2/run.sh e2e dashboard --run-id v2_15seed"],
        ],
        [1800, 7560],
    )

    document.add_heading("5. AI 도구에 읽힐 파일", level=1)
    _add_table(
        document,
        ["도구", "먼저 읽힐 파일", "역할별 추가 파일"],
        [
            ["Claude", "v2/CLAUDE.md", "v2/ai_skills/<role>/SKILL.md"],
            ["Gemini", "v2/GEMINI.md", "v2/ai_skills/<role>/SKILL.md"],
            ["Cursor", "v2/ai_skills/common_project_rules.md", "v2/ai_skills/<role>/SKILL.md"],
            ["Antigravity", "v2/ai_skills/common_project_rules.md", "v2/ai_skills/hatespeech-v2-e2e/SKILL.md"],
            ["Codex", "v2/ai_skills/common_project_rules.md", "필요 시 hatespeech-v2-* skill 폴더"],
        ],
        [1500, 4100, 3760],
    )

    document.add_heading("6. 완료 보고 양식", level=1)
    _add_callout(
        document,
        "반드시 이 형식으로 보고",
        "[v2 agent handoff]\nRole:\nFiles changed:\nCommands run:\nArtifacts created/updated:\nValidation passed:\nKnown limitations:\nNext owner:",
    )

    document.add_heading("7. 즉시 업무 도입 판정", level=1)
    _add_bullets(
        document,
        [
            "문서 목록은 충분하지만, 목록만 읽으면 업무 행동으로 전환하는 데 시간이 걸린다.",
            "이 Word 문서는 D0 업무 하달 직후 팀원이 바로 자기 역할을 시작하도록 만든 실행형 보조 문서다.",
            "팀장은 17_korean_reading_file_index.md와 이 Word를 함께 공유하면 된다.",
        ],
    )

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
