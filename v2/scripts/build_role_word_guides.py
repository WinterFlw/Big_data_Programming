"""Build role-specific Korean DOCX guides for the five-person v2 team.

The team split is intentionally based on research deliverables, not on a
low-level source-code dependency graph.  Each generated guide answers the
questions a teammate asks on day one:

* What am I responsible for?
* Which v2 files do I read first?
* Which commands do I run?
* Which artifacts prove that my part is done?
* How do I report progress back to the team lead?

The documents are generated deterministically so that future edits can be made
in this one script and re-run without manually touching multiple Word files.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_DIR = BASE_DIR / "docs" / "role_guides"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WARNING = "FFF2CC"
SUCCESS = "E2F0D9"


@dataclass(frozen=True)
class RoleGuide:
    """All content needed to generate one role handoff document."""

    number: int
    slug: str
    title: str
    subtitle: str
    owner_mission: str
    one_line: str
    responsibilities: list[str]
    read_first: list[str]
    code_scope: list[str]
    commands: list[tuple[str, str]]
    artifacts: list[str]
    done_checks: list[str]
    risks: list[str]
    ai_prompt: str
    special_sections: list[tuple[str, list[str]]]


def _set_cell_shading(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    """Set explicit DXA cell padding so Korean text does not touch borders."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, widths: list[int]) -> None:
    """Use fixed Word table geometry instead of renderer-dependent autofit."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for existing in list(grid):
        grid.remove(existing)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 120, line: int = 280) -> None:
    """Patch paragraph spacing in OOXML so list items remain readable."""
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def _style_document(document: Document, footer_label: str) -> None:
    """Apply the standard business brief style used by the other v2 DOCX files."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(footer_label)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def _add_title(document: Document, guide: RoleGuide) -> None:
    title = document.add_paragraph(style="Title")
    title.add_run(guide.title)
    subtitle = document.add_paragraph()
    subtitle.add_run(guide.subtitle).bold = True
    meta = document.add_paragraph()
    meta.add_run("기준 workspace: ").bold = True
    meta.add_run("v2/   ")
    meta.add_run("Run ID: ").bold = True
    meta.add_run("v2_15seed   ")
    meta.add_run("산출물 루트: ").bold = True
    meta.add_run("v2/outputs/experiments/v2_15seed/")


def _add_callout(document: Document, title: str, body: str, fill: str = CALLOUT) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _set_table_width(table, [9360])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    paragraph.add_run(f"\n{body}")
    document.add_paragraph()


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        _set_paragraph_spacing(paragraph, after=80, line=280)


def _add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)
        _set_paragraph_spacing(paragraph, after=80, line=280)


def _add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        _set_cell_shading(cell, LIGHT_GRAY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    _set_table_width(table, widths)
    document.add_paragraph()


def _add_command_table(document: Document, commands: list[tuple[str, str]]) -> None:
    _add_table(document, ["목적", "명령"], [[purpose, command] for purpose, command in commands], [2100, 7260])


def _add_report_template(document: Document) -> None:
    _add_callout(
        document,
        "공통 보고 양식",
        "맡은 역할 / 확인한 파일 / 실행한 명령 / 생성된 산출물 / 발견한 문제 / 수정한 내용 / 아직 불확실한 부분 / 다음 사람이 이어받아야 할 것",
        fill=SUCCESS,
    )


def _build_one(guide: RoleGuide) -> Path:
    document = Document()
    _style_document(document, f"HateSpeachStudy v2 role guide {guide.number}: {guide.slug}")
    _add_title(document, guide)

    _add_callout(document, "한 줄 책임", guide.one_line)

    document.add_heading("1. 역할 목적", level=1)
    document.add_paragraph(guide.owner_mission)

    document.add_heading("2. 담당 업무", level=1)
    _add_bullets(document, guide.responsibilities)

    document.add_heading("3. 먼저 읽을 파일", level=1)
    _add_bullets(document, guide.read_first)

    document.add_heading("4. 책임 코드/산출물 범위", level=1)
    _add_table(
        document,
        ["구분", "범위"],
        [["코드/설정", "\n".join(guide.code_scope)], ["완료 산출물", "\n".join(guide.artifacts)]],
        [1800, 7560],
    )

    document.add_heading("5. 첫 실행 명령", level=1)
    _add_command_table(document, guide.commands)

    for heading, items in guide.special_sections:
        document.add_heading(heading, level=1)
        _add_bullets(document, items)

    document.add_heading("6. 완료 기준", level=1)
    _add_bullets(document, guide.done_checks)

    document.add_heading("7. 주의할 위험", level=1)
    _add_bullets(document, guide.risks)

    document.add_heading("8. AI 도구에 줄 지시문", level=1)
    _add_callout(document, "복붙 프롬프트", guide.ai_prompt)

    document.add_heading("9. 보고 방식", level=1)
    _add_report_template(document)

    path = OUTPUT_DIR / f"{guide.number:02d}_{guide.slug}.docx"
    document.save(path)
    return path


GUIDES: list[RoleGuide] = [
    RoleGuide(
        number=1,
        slug="code_review_pipeline_validation",
        title="1번 E2E Gate 총괄 업무지시서",
        subtitle="전체 코드리뷰가 아니라 stage 연결부와 full run GO/STOP을 판단하는 역할",
        owner_mission=(
            "이 역할은 전체 코드를 다 외우거나 모든 파트를 리뷰하는 사람이 아니다. "
            "각 담당자가 자기 파트 코드를 1차 리뷰한 결과를 취합하고, limited server run 전에 "
            "run_id, output path, stage 연결, smoke gate가 문서 계약과 맞는지 판단한다."
        ),
        one_line="각 담당자의 1차 리뷰 결과를 모아 full benchmark GO/STOP을 판단한다.",
        responsibilities=[
            "2~5번 담당자의 1차 리뷰 보고를 취합한다.",
            "v2/run.sh, CLI, manifest, paths, artifacts, scripts gate만 깊게 본다.",
            "dry-run, aggregate, xai-bundle, report, dashboard가 같은 run_id/output root를 공유하는지 확인한다.",
            "full 120 run 전에 P0/P1 문제와 나중에 고쳐도 되는 문제를 구분한다.",
            "GO/STOP 판단을 팀장에게 보고한다.",
        ],
        read_first=[
            "v2/README.md",
            "v2/docs/02_e2e_pipeline.md",
            "v2/docs/06_execution_runbook.md",
            "v2/docs/07_output_and_report_contract.md",
            "v2/docs/15_runtime_code_validation_matrix.md",
            "v2/docs/20_role_file_review_matrix.md",
        ],
        code_scope=[
            "v2/run.sh",
            "v2/pipeline/cli.py",
            "v2/pipeline/runner.py",
            "v2/pipeline/manifest.py",
            "v2/pipeline/paths.py",
            "v2/pipeline/schema.py",
            "v2/pipeline/artifacts.py",
            "v2/configs/v2_15seed.json",
            "v2/scripts/daily.sh",
            "v2/scripts/gate_check.py",
        ],
        commands=[
            ("문법 확인", "python3 -m compileall v2/runtime v2/pipeline v2/scripts"),
            ("config 확인", "python3 -m json.tool v2/configs/v2_15seed.json >/tmp/v2_config_check.json"),
            ("dry-run", "PYTHON_BIN=python3 ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B,D_B --seeds 42 --dry-run"),
            ("집계 smoke", "PYTHON_BIN=python3 ./v2/run.sh e2e aggregate --run-id v2_15seed"),
            ("XAI bundle smoke", "PYTHON_BIN=python3 ./v2/run.sh e2e xai-bundle --run-id v2_15seed"),
            ("report/dashboard smoke", "PYTHON_BIN=python3 ./v2/run.sh e2e report --run-id v2_15seed && PYTHON_BIN=python3 ./v2/run.sh e2e dashboard --run-id v2_15seed"),
            ("gate check", "python3 v2/scripts/gate_check.py --run-id v2_15seed --skip-sample-check"),
        ],
        artifacts=[
            "v2/outputs/experiments/v2_15seed/execution_status.csv",
            "v2/outputs/experiments/v2_15seed/benchmark/*.csv",
            "v2/outputs/experiments/v2_15seed/xai/evidence_bundle/*.json",
            "v2/outputs/experiments/v2_15seed/reports/final_report.md",
            "v2/outputs/experiments/v2_15seed/dashboard/index.html",
            "검증 결과 체크리스트",
        ],
        done_checks=[
            "dry-run이 2개 selected unit을 생성한다.",
            "aggregate/report/dashboard stage가 빈 결과 또는 smoke 결과에서 죽지 않는다.",
            "output root가 v2_15seed 아래로 고정된다.",
            "full 120 run 전에 막아야 할 P0/P1 문제가 목록화된다.",
        ],
        risks=[
            "실제 학습 성공과 dry-run 성공을 혼동하면 안 된다.",
            "2~5번의 파트별 코드리뷰를 대신 떠안으면 병목이 된다.",
            "v1 archive 문서를 실행 기준으로 삼으면 안 된다.",
            "경로 문제는 서버에서 바로 시간 손실로 이어지므로 사소하게 보지 않는다.",
        ],
        ai_prompt=(
            "너는 HateSpeachStudy v2의 E2E Gate 총괄 담당자다. "
            "전체 코드리뷰를 떠안지 말고 v2/run.sh, v2/pipeline/cli.py, runner.py, "
            "manifest.py, paths.py, artifacts.py, v2/scripts/daily.sh, gate_check.py를 중심으로 "
            "run_id, output path, stage 연결, smoke/full gate가 맞는지 검토해라. "
            "2~5번 담당자의 1차 리뷰 결과를 취합해 GO/STOP 판단표를 만들어라."
        ),
        special_sections=[
            (
                "1번이 직접 리뷰하지 않는 파일",
                [
                    "학습 세부 로직: 2번이 v2/pipeline/training_adapter.py, v2/runtime/experiment_core.py, v2/runtime/run_experiments.py, v2/runtime/utils.py를 1차 리뷰한다.",
                    "통계 세부 로직: 3번이 v2/pipeline/statistics.py와 v2/pipeline/schema.py를 1차 리뷰한다.",
                    "XAI 세부 로직: 4번이 v2/pipeline/xai.py, v2/pipeline/xai_sampling.py, v2/pipeline/xai_bundle.py, v2/runtime/experiment_xai.py를 1차 리뷰한다.",
                    "보고서/대시보드: 5번이 v2/pipeline/reporting.py, v2/runtime/dashboard_app.py, v2/runtime/experiment_dashboard.py를 1차 리뷰한다.",
                ],
            ),
            (
                "서버 실행 전 금지선",
                [
                    "A_B seed 42 실제 smoke 전에는 full benchmark를 시작하지 않는다.",
                    "A_B/D_B seed 42 paired smoke 전에는 통계 결과를 확정하지 않는다.",
                    "report/dashboard가 생성되지 않으면 발표자료 담당에게 결과를 넘기지 않는다.",
                ],
            )
        ],
    ),
    RoleGuide(
        number=2,
        slug="training_execution_experiment_management",
        title="2번 학습 실행 / 실험 관리 업무지시서",
        subtitle="서버에서 seed와 condition을 통제해 실제 학습 결과를 만드는 역할",
        owner_mission=(
            "이 역할은 연구 결과의 원천 데이터를 만드는 사람이다. 15 seed x 8 condition "
            "실행을 관리하고, 실패 run을 추적하며, 같은 seed가 모든 조건에 공통으로 적용되도록 보장한다."
        ),
        one_line="서버 시간을 실제 학습 결과로 바꾸고, 실패 run을 재현 가능하게 관리한다.",
        responsibilities=[
            "GPU 서버 환경에서 v2 학습 의존성, 데이터 경로, PYTHON_BIN을 확인한다.",
            "A_B seed 42 단일 smoke를 먼저 실행한다.",
            "A_B/D_B seed 42 paired smoke로 비교 가능한 결과가 생성되는지 확인한다.",
            "15 seed 전체 실행 시 completed_runs.csv와 failed_runs.csv를 관리한다.",
            "실패한 run은 전체 재실행이 아니라 실패 단위만 resume/retry한다.",
        ],
        read_first=[
            "v2/docs/06_execution_runbook.md",
            "v2/docs/11_team_tasking_and_server_run_plan.md",
            "v2/docs/07_output_and_report_contract.md",
            "v2/docs/15_runtime_code_validation_matrix.md",
            "v2/docs/20_role_file_review_matrix.md",
            "v2/docs/agent_tasks/01_benchmark_agent.md",
        ],
        code_scope=[
            "v2/pipeline/training_adapter.py",
            "v2/pipeline/artifacts.py",
            "v2/runtime/experiment_core.py",
            "v2/runtime/run_experiments.py",
            "v2/runtime/utils.py",
            "v2/runtime/requirements.txt",
            "v2/configs/v2_15seed.json",
        ],
        commands=[
            ("상태 확인", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e status --run-id v2_15seed"),
            ("단일 학습 smoke", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B --seeds 42 --execute --resume"),
            ("paired smoke", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e benchmark --run-id v2_15seed --conditions A_B,D_B --seeds 42 --execute --resume"),
            ("집계 확인", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e aggregate --run-id v2_15seed"),
            ("전체 실행", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e benchmark --run-id v2_15seed --execute --resume"),
        ],
        artifacts=[
            "condition x seed별 metrics.json",
            "condition x seed별 history.csv",
            "condition x seed별 predictions.csv",
            "condition x seed별 run_config.json",
            "condition x seed별 checkpoint",
            "completed_runs.csv / failed_runs.csv / execution_status.csv",
        ],
        done_checks=[
            "A_B seed 42 smoke가 metrics/history/predictions/checkpoint를 만든다.",
            "A_B와 D_B가 같은 seed 42로 비교 가능한 상태가 된다.",
            "실패 run 목록이 누락 없이 기록된다.",
            "15 seed 전체 실행 후 completed count와 planned count의 차이를 설명할 수 있다.",
        ],
        risks=[
            "seed가 조건별로 다르면 paired design이 깨진다.",
            "checkpoint만 있고 predictions/metrics가 없으면 분석 담당이 쓸 수 없다.",
            "서버에서 경로를 바꿔 실행하면 report/dashboard가 결과를 못 찾을 수 있다.",
            "실패 run을 덮어쓰기 방식으로 재실행하면 원인 추적이 어려워진다.",
        ],
        ai_prompt=(
            "너는 HateSpeachStudy v2의 학습 실행/실험 관리 담당자다. "
            "A_B seed 42 smoke부터 시작해 condition, seed, run_id, output artifact가 "
            "정확히 남는지 확인해라. 서버 실행 명령, 실패 run, 재실행 기준을 표로 정리해라."
        ),
        special_sections=[
            (
                "실행 순서",
                [
                    "1단계: config와 데이터 경로 확인",
                    "2단계: A_B seed 42 단일 smoke",
                    "3단계: A_B/D_B seed 42 paired smoke",
                    "4단계: aggregate/report/dashboard가 smoke 결과를 읽는지 확인",
                    "5단계: 15 seed full benchmark",
                ],
            )
        ],
    ),
    RoleGuide(
        number=3,
        slug="result_analysis_statistics",
        title="3번 결과 분석 / 통계 해석 업무지시서",
        subtitle="15 seed 결과가 통계적으로 어떤 의미인지 설명하는 역할",
        owner_mission=(
            "이 역할은 모델이 정말 나아졌는지 판단한다. 단일 최고 성능이 아니라 15 seed의 "
            "분포, 핵심 paired difference, confidence interval, effect size를 함께 본다. "
            "Holm과 ANOVA는 여러 조건을 한꺼번에 볼 때의 보조 분석으로만 다룬다."
        ),
        one_line="15 seed 성능표를 학부 수준에 맞는 결론으로 바꾸되, p-value 하나에 기대지 않는다.",
        responsibilities=[
            "benchmark_runs.csv와 benchmark_summary.csv가 완성됐는지 확인한다.",
            "핵심 비교인 A_B vs D_B를 같은 seed 기반 paired t-test로 검증한다.",
            "mean difference, 95% CI, effect size, raw paired p-value를 함께 해석한다.",
            "여러 비교와 ANOVA는 보조/부록 분석으로 낮춰서 정리한다.",
            "발표자료에 들어갈 통계 문장과 표를 만든다.",
        ],
        read_first=[
            "v2/docs/03_validation_and_statistics.md",
            "v2/docs/07_output_and_report_contract.md",
            "v2/docs/08_xai_report_template.md",
            "v2/docs/20_role_file_review_matrix.md",
            "v2/docs/agent_tasks/02_statistics_agent.md",
            "v2/docs/agent_tasks/23_role_responsibilities.md",
        ],
        code_scope=[
            "v2/pipeline/statistics.py",
            "v2/pipeline/schema.py",
            "v2/docs/03_validation_and_statistics.md",
            "v2/docs/07_output_and_report_contract.md",
            "v2/outputs/experiments/v2_15seed/benchmark/*.csv",
        ],
        commands=[
            ("집계 실행", "PYTHON_BIN=python3 ./v2/run.sh e2e aggregate --run-id v2_15seed"),
            ("summary 확인", "column -t -s, v2/outputs/experiments/v2_15seed/benchmark/benchmark_summary.csv | head"),
            ("핵심 paired 확인", "column -t -s, v2/outputs/experiments/v2_15seed/benchmark/paired_tests_holm.csv | head"),
            ("보조 ANOVA 확인", "ls v2/outputs/experiments/v2_15seed/benchmark/anova_*.csv"),
            ("report 반영 확인", "PYTHON_BIN=python3 ./v2/run.sh e2e report --run-id v2_15seed"),
        ],
        artifacts=[
            "benchmark_summary.csv",
            "paired_tests.csv",
            "paired_tests_holm.csv (adjusted p-value는 보조 확인용)",
            "anova_2way_bert.csv / anova_2way_roberta.csv / anova_3way.csv (부록용)",
            "발표용 통계표",
            "결과 해석 문장",
        ],
        done_checks=[
            "각 condition의 n_seeds가 실제 완료 seed 수와 일치한다.",
            "A_B vs D_B paired test의 n_pairs가 비교 가능한 같은 seed 개수를 반영한다.",
            "최종 결론은 mean difference, CI, effect size, paired p-value를 함께 보고 쓴다.",
            "효과크기와 CI가 결론 문장에 함께 들어간다.",
        ],
        risks=[
            "best seed만 보고 결론을 쓰면 안 된다.",
            "Holm이나 ANOVA를 발표의 중심에 놓으면 학부 프로젝트 치고 설명 부담이 커진다.",
            "n_pairs가 너무 작으면 p-value가 의미 있는 근거가 되기 어렵다.",
            "통계적 유의성과 실제 효과크기는 다를 수 있다.",
        ],
        ai_prompt=(
            "너는 HateSpeachStudy v2의 결과 분석/통계 해석 담당자다. "
            "15 seed mean/std, A_B vs D_B same-seed paired t-test, "
            "CI, effect size를 중심으로 성능 개선 여부를 과장 없이 설명해라. "
            "Holm-adjusted p-value와 ANOVA는 여러 조건을 보여줄 때의 보조/부록 분석으로만 써라."
        ),
        special_sections=[
            (
                "우리 수준의 분석 기준",
                [
                    "메인은 15 seed condition별 Macro F1 mean ± std다.",
                    "핵심 검정은 A_B baseline과 D_B 제안 조건의 same-seed paired t-test다.",
                    "p-value만 말하지 말고 평균 차이, 95% CI, effect size를 함께 말한다.",
                    "Holm 보정, ANOVA, bootstrap은 계산할 수 있지만 발표 본문에서는 보조/부록으로 낮춘다.",
                    "XAI는 통계 결론을 증명하는 도구가 아니라 결과를 해석하는 보조 근거다.",
                ],
            ),
            (
                "Holm 보정은 어디에 쓰는가?",
                [
                    "Holm-Bonferroni 보정은 여러 개의 통계 검정을 동시에 할 때 false positive를 줄이는 다중비교 보정 방법이다.",
                    "예를 들어 조건 비교를 1개만 하면 p < 0.05 기준이 비교적 단순하지만, 7개 비교를 동시에 하면 우연히 하나쯤 유의하게 나올 확률이 커진다.",
                    "그래서 여러 조건 비교표를 같이 보여줄 때만 adjusted p-value를 보조로 확인한다.",
                    "발표 본문에서는 Holm을 길게 설명하지 말고 '다중 비교 과대해석 방지용 보조 확인' 정도로만 쓴다.",
                    "우리 결론의 중심은 A_B vs D_B paired t-test, 평균 차이, CI, effect size다.",
                ],
            ),
            (
                "발표용 문장 예시",
                [
                    "본 연구는 15개 random seed 반복 실험으로 단일 실행 결과 의존을 줄였다.",
                    "핵심 비교는 baseline A_B와 제안 조건 D_B의 같은 seed 기반 paired t-test로 수행했다.",
                    "p-value뿐 아니라 평균 차이, 95% CI, effect size를 함께 고려했다.",
                    "여러 조건을 동시에 비교하는 표에서는 adjusted p-value를 보조적으로 확인했다.",
                ],
            ),
        ],
    ),
    RoleGuide(
        number=4,
        slug="xai_explanation_evidence_bundle",
        title="4번 XAI 설명 / Evidence Bundle 업무지시서",
        subtitle="성능 차이를 사람이 이해할 수 있는 설명 근거로 정리하는 역할",
        owner_mission=(
            "이 역할은 XAI 그림 몇 장을 고르는 사람이 아니라, explanation artifact를 "
            "검증 가능한 evidence bundle로 묶는 사람이다. XAI는 성능 개선의 절대 증명이 아니라 "
            "모델 판단 패턴이 연구 가설과 일관적인지 확인하는 사후 근거다."
        ),
        one_line="XAI를 예쁜 시각화가 아니라 report/dashboard가 읽는 근거 번들로 만든다.",
        responsibilities=[
            "Primary XAI에서 A_B와 D_B를 같은 sample, 같은 seed 기준으로 비교한다.",
            "Deep XAI에서 median-performing checkpoint 기반 대표 case를 정리한다.",
            "Ablation XAI에서 8조건의 설명 경향을 가볍게 비교한다.",
            "xai_claims.json과 xai_dashboard_bundle.json의 claim/source 연결을 확인한다.",
            "XAI 결과를 성능의 인과 증명처럼 과장하지 않고 보조 근거로 설명한다.",
        ],
        read_first=[
            "v2/docs/04_xai_protocol.md",
            "v2/docs/07_output_and_report_contract.md",
            "v2/docs/08_xai_report_template.md",
            "v2/docs/20_role_file_review_matrix.md",
            "v2/docs/agent_tasks/03_xai_agent.md",
            "v2/docs/agent_tasks/09_e2e_xai_evidence_bundle_agent.md",
        ],
        code_scope=[
            "v2/pipeline/xai.py",
            "v2/pipeline/xai_sampling.py",
            "v2/pipeline/xai_bundle.py",
            "v2/runtime/experiment_xai.py",
            "v2/outputs/experiments/v2_15seed/xai/",
        ],
        commands=[
            ("primary XAI", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e xai-primary --run-id v2_15seed --resume"),
            ("deep XAI", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e xai-deep --run-id v2_15seed --resume"),
            ("ablation XAI", "PYTHON_BIN=/path/to/venv/bin/python ./v2/run.sh e2e xai-ablation --run-id v2_15seed --resume"),
            ("bundle 생성", "PYTHON_BIN=python3 ./v2/run.sh e2e xai-bundle --run-id v2_15seed"),
            ("JSON 검증", "python3 -m json.tool v2/outputs/experiments/v2_15seed/xai/evidence_bundle/xai_claims.json >/dev/null"),
        ],
        artifacts=[
            "xai/primary/seed_level_metrics.csv",
            "xai/primary/paired_xai_tests.csv",
            "xai/primary/seed_stability.csv",
            "xai/deep/case_summary.csv",
            "xai/ablation/xai_ablation_metrics.csv",
            "xai/evidence_bundle/xai_claims.json",
            "xai/evidence_bundle/xai_dashboard_bundle.json",
            "xai/evidence_bundle/token_attributions.jsonl",
        ],
        done_checks=[
            "primary sample set이 seed마다 바뀌지 않는다.",
            "xai_claims.json의 각 claim에 source_artifacts가 붙어 있다.",
            "report/dashboard가 raw XAI case보다 evidence bundle을 우선 입력으로 쓴다.",
            "대표 case가 전체 통계 결론을 대신하지 않는다.",
        ],
        risks=[
            "XAI를 모델이 맞았다는 증명으로 표현하면 안 된다.",
            "성공 case만 골라 발표하면 cherry-picking이 된다.",
            "sample이 seed마다 바뀌면 explanation stability를 주장할 수 없다.",
            "claim에 source artifact가 없으면 발표 근거로 쓰지 않는다.",
        ],
        ai_prompt=(
            "너는 HateSpeachStudy v2의 XAI explanation/evidence bundle 담당자다. "
            "XAI를 사후 검증 근거로 다루고, xai_claims.json과 xai_dashboard_bundle.json이 "
            "어떤 통계/XAI artifact에서 왔는지 확인해라. 대표 case는 과장하지 말고, "
            "seed stability와 source_artifacts를 기준으로 보고해라."
        ),
        special_sections=[
            (
                "XAI 해석 기준",
                [
                    "Attribution alignment: 중요 토큰이 인간 rationale과 겹치는가?",
                    "Faithfulness: 중요 토큰을 제거하거나 가렸을 때 예측이 흔들리는가?",
                    "Context learning: target/source/context subgroup에서 설명 패턴이 일관적인가?",
                    "Plausibility: 사람이 보기에 설명이 납득 가능한가?",
                    "Seed stability: seed가 달라도 중요한 토큰과 설명 순위가 유지되는가?",
                ],
            )
        ],
    ),
    RoleGuide(
        number=5,
        slug="presentation_report_final_integration",
        title="5번 발표자료 / 최종 보고서 제작 업무지시서",
        subtitle="코드, 학습, 통계, XAI 결과를 최종 제출물로 묶는 역할",
        owner_mission=(
            "이 역할은 결과를 예쁘게 포장하는 역할이 아니라, 연구 질문부터 실험 설계, "
            "성능 결과, 통계 근거, XAI 근거, 한계까지 하나의 발표 흐름으로 엮는 최종 통합 담당자다."
        ),
        one_line="나머지 네 역할의 산출물을 발표와 보고서로 바꾼다.",
        responsibilities=[
            "최종 보고서가 benchmark_summary, paired_tests_holm, xai evidence bundle을 읽는지 확인한다.",
            "PPT 흐름을 연구 질문, 방법, 실험 설계, 결과, XAI, 한계, 결론으로 구성한다.",
            "통계 담당의 문장과 XAI 담당의 문장을 서로 충돌하지 않게 합친다.",
            "dashboard/index.html과 final_report.docx의 위치를 팀에 공유한다.",
            "발표에서 과장된 인과 주장, p-value 단독 결론, Holm/ANOVA 과잉 설명을 제거한다.",
        ],
        read_first=[
            "v2/docs/01_model_definition.md",
            "v2/docs/02_e2e_pipeline.md",
            "v2/docs/03_validation_and_statistics.md",
            "v2/docs/04_xai_protocol.md",
            "v2/docs/20_role_file_review_matrix.md",
            "v2/docs/07_output_and_report_contract.md",
            "v2/docs/08_xai_report_template.md",
            "v1/docs/발표_와꾸_v2.md",
        ],
        code_scope=[
            "v2/pipeline/reporting.py",
            "v2/runtime/dashboard_app.py",
            "v2/runtime/experiment_dashboard.py",
            "v2/outputs/experiments/v2_15seed/reports/",
            "v2/outputs/experiments/v2_15seed/dashboard/",
            "v2/docs/",
        ],
        commands=[
            ("report 생성", "PYTHON_BIN=python3 ./v2/run.sh e2e report --run-id v2_15seed"),
            ("dashboard 생성", "PYTHON_BIN=python3 ./v2/run.sh e2e dashboard --run-id v2_15seed"),
            ("report 확인", "head -80 v2/outputs/experiments/v2_15seed/reports/final_report.md"),
            ("dashboard 위치 확인", "ls -la v2/outputs/experiments/v2_15seed/dashboard/"),
            ("XAI bundle JSON 확인", "python3 -m json.tool v2/outputs/experiments/v2_15seed/xai/evidence_bundle/xai_dashboard_bundle.json >/dev/null"),
        ],
        artifacts=[
            "final_report.md",
            "final_report.docx",
            "dashboard/index.html",
            "PPT 초안",
            "발표 스크립트",
            "최종 제출물 체크리스트",
        ],
        done_checks=[
            "보고서가 통계 결과와 XAI bundle을 모두 반영한다.",
            "PPT에 15 seed, 핵심 paired t-test, mean difference, CI, effect size 설명이 들어간다.",
            "Holm/ANOVA는 필요할 때만 부록 또는 보조 검증으로 짧게 언급한다.",
            "XAI는 보조 근거로 설명되고 인과 증명처럼 쓰이지 않는다.",
            "최종 산출물 위치와 재현 명령이 마지막 슬라이드 또는 부록에 적힌다.",
        ],
        risks=[
            "성능 차이만 강조하면 v2의 XAI evidence bundle 강점이 사라진다.",
            "XAI case를 너무 많이 넣으면 발표 흐름이 산만해진다.",
            "Holm과 ANOVA를 길게 설명하면 발표가 통계 수업처럼 보일 수 있다.",
            "최종 파일 위치를 모르면 제출 직전에 산출물을 다시 찾게 된다.",
        ],
        ai_prompt=(
            "너는 HateSpeachStudy v2의 발표자료/최종 보고서 제작 담당자다. "
            "통계 결과와 XAI evidence bundle을 입력으로 받아 연구 질문, 방법, 결과, 해석, "
            "한계, 결론 흐름으로 PPT와 보고서 문장을 구성해라. p-value 단독 결론과 "
            "XAI 인과 과장을 제거해라."
        ),
        special_sections=[
            (
                "추천 발표 흐름",
                [
                    "문제 정의: 혐오표현 탐지에서 문맥과 rationale이 왜 중요한가",
                    "모델 정의: baseline, rationale-aware attention, VADER feature, 8조건 ablation",
                    "실험 설계: 15 seed, same-seed paired design, output contract",
                    "성능 결과: mean/std, 핵심 paired t-test, mean difference, CI, effect size",
                    "XAI 결과: seed stability, representative cases, evidence bundle",
                    "한계: XAI는 보조 근거이며 인과 증명이 아님",
                    "결론: 정확도 차이와 설명 가능성 산출물을 함께 제시",
                ],
            )
        ],
    ),
]


def build_all() -> list[Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return [_build_one(guide) for guide in GUIDES]


def main() -> None:
    for path in build_all():
        print(path)


if __name__ == "__main__":
    main()
